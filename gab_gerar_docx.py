# gerar_documento_GAB.py
from docx import Document
from docx.shared import Cm, Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from gab_prazo_pag import construir_primeira_pagina
from gab_tempo_pag import construir_segunda_pagina
from config import PASTA_SAIDA
import os


def gerar_documento_gab():
    pasta_gab = os.path.join(PASTA_SAIDA, "GAB")
    doc = Document()

    # --- Configuração de margens ---
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # --- Fonte padrão ---
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # --- Caminhos das imagens ---
    imagens = {
        # PRAZO PADRÃO
        "card_prazo": f"{pasta_gab}/GAB_Prazo_Cards.png",
        "top10_prazo": f"{pasta_gab}/GAB_Prazo_Top10.png",
        "top3_prazo": f"{pasta_gab}/GAB_Prazo_Top3.png",
        "grafico_prazo": f"{pasta_gab}/GAB_Prazo_Grafico_6meses.png",

        # TEMPO PADRÃO
        "card_tempo": f"{pasta_gab}/GAB_Tempo_Cards.png",
        "top10_tempo": f"{pasta_gab}/GAB_Tempo_Top10.png",
        "top3_tempo": f"{pasta_gab}/GAB_Tempo_Top3.png",
        "grafico_tempo": f"{pasta_gab}/GAB_Tempo_Grafico_6meses.png",

        # NOVO GRÁFICO GERAL POR CONCESSIONÁRIA
        "grafico_conc": f"{pasta_gab}/GAB_Concessionarias_Prazo.png"
    }

    # --- Página 1: Prazo Padrão ---
    construir_primeira_pagina(doc, "Report Cesta de Serviços – GAB", "Outubro", 2025, imagens)

    # --- Quebra de página ---
    doc.add_page_break()

    # --- Página 2: Tempo Padrão ---
    construir_segunda_pagina(doc, "Report Cesta de Serviços – GAB", "Outubro", 2025, imagens)

    # ============================================================
    #  🚀 INSERIR NOVO GRÁFICO — por concessionária (sem pular página)
    # ============================================================
    if os.path.exists(imagens["grafico_conc"]):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run()
        run.add_picture(imagens["grafico_conc"], width=Inches(6))
        doc.add_paragraph("\n")

    # --- Salvar arquivo ---
    caminho_doc = os.path.join(pasta_gab, "Report_GAB.docx")
    doc.save(caminho_doc)

    print(f"📄 Documento consolidado GAB salvo em: {caminho_doc}")


if __name__ == "__main__":
    gerar_documento_gab()
