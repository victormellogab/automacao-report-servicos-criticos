import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import plotly.graph_objects as go

# ----------- CONFIGURAÇÃO -----------
PASTA_SAIDA = r'H:\COMERCIAL\MEDIÇÃO\ARQUIVO MORTO\BERNARDO\GSC\Automação Report'
os.makedirs(PASTA_SAIDA, exist_ok=True)

CONCESSIONARIAS = ['CAC']

CONCESSIONARIA_NOME = {
    'CAC': 'Concessionária Águas das Agulhas Negras',
}

MES = "Outubro"
ANO = 2025

IMAGENS = {
    'CAC': {
        "card_prazo": os.path.join(PASTA_SAIDA, "CAC_Resumo_Setembro.png"),
        "top10_prazo": os.path.join(PASTA_SAIDA, "CAC_Top10_Setembro.png"),
        "top3_prazo": os.path.join(PASTA_SAIDA, "CAC_Top3_3Meses.png"),
        "grafico_prazo": os.path.join(PASTA_SAIDA, "CAC_Evolucao_Abr-Set_2025.png"),
        "card_tempo": None,
        "top10_tempo": None,
        "top3_tempo": None,
        "grafico_tempo": None,
        "tabela_prioritarios": None
    }
}

# Caminho do Excel da tabela de serviços prioritários
CAMINHO_TABELA = r"C:\Users\victor.mello\OneDrive - Grupo Aguas do Brasil\Área de Trabalho\Script Geração Report Serviços\Dados\Cesta de Serviços.xlsx"
IMG_TABELA = os.path.join(PASTA_SAIDA, "Tabela_Prioritarios.png")

# ----------- FUNÇÃO PARA GERAR IMAGEM DA TABELA -----------
def gerar_imagem_tabela(excel_path, caminho_saida):
    df = pd.read_excel(excel_path)

    # Ajustar largura das colunas proporcional ao tamanho máximo do conteúdo
    max_lens = [max(df[col].astype(str).map(len).max(), len(col)) for col in df.columns]
    
    # Diminuir a largura da coluna "Serviços" (primeira coluna)
    max_lens[0] = int(max_lens[0] * 0.16)  # 60% do tamanho proporcional
    
    total_len = sum(max_lens)
    column_widths = [length / total_len for length in max_lens]

    fig = go.Figure(data=[go.Table(
        columnwidth=column_widths,
        header=dict(
            values=list(df.columns),
            fill_color='darkblue',
            font=dict(color='white', size=14),
            align='center'
        ),
        cells=dict(
            values=[df[col] for col in df.columns],
            fill_color='white',
            font=dict(color='black', size=12),
            align=['left' if col=='Serviços' else 'center' for col in df.columns]
        )
    )])

    # Altura dinâmica: 40px por linha + 40px para o cabeçalho
    altura = 21 + 21 * len(df)
    
    # Layout com altura calculada e sem margens extras
    fig.update_layout(
        margin=dict(l=0, r=0, t=0, b=0),
        width=None,
        height=altura
    )

    fig.write_image(caminho_saida, scale=3)

# ----------- FUNÇÃO PARA ADICIONAR IMAGEM OU TEXTO -----------
def add_imagem_ou_texto(doc, texto, caminho_imagem, tamanho_pct=100, align='center'):
    p = doc.add_paragraph()
    
    # Definir alinhamento
    if align == 'center':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    elif align == 'left':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    elif align == 'right':
        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    else:
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    if caminho_imagem and os.path.exists(caminho_imagem):
        largura_inch = Inches(6) * (tamanho_pct / 100)
        run = p.add_run()
        run.add_picture(caminho_imagem, width=largura_inch)
    else:
        run = p.add_run(texto)
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        run.font.size = Pt(12)
        run.font.bold = False
        run.font.color.rgb = RGBColor(0,0,0)

# ----------- FUNÇÃO DE CRIAÇÃO DO DOCUMENTO -----------
def gerar_documento_report(concessionaria, pasta_saida, imagens):
    # Gera a imagem da tabela
    gerar_imagem_tabela(CAMINHO_TABELA, IMG_TABELA)
    
    doc = Document()

    # ---------- CONFIGURAR MARGENS ----------
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ---------- CONFIGURAR FONTE PADRÃO ----------
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Configurar estilos de títulos
    for level in range(1, 10):
        try:
            s = doc.styles[f'Heading {level}']
            s.font.name = 'Calibri'
            s.font.size = Pt(12)
            s.font.bold = False
            s.font.color.rgb = RGBColor(0,0,0)
            s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        except KeyError:
            pass

    # Nome completo da concessionária
    nome_conc = CONCESSIONARIA_NOME.get(concessionaria, concessionaria)

    # Título
    p = doc.add_paragraph()
    run = p.add_run(nome_conc)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(21)
    run.add_break()  # quebra de linha

    # Subtítulo
    run2 = p.add_run(f"Report Cesta de Serviços – {MES} {ANO}")
    run2.font.name = 'Calibri'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run2.font.size = Pt(12)

    # Seções principais
    seccoes = [
        ("Indicador de Serviços no Prazo Padrão", "card_prazo", 70),
        ("Top 10 Serviços com Menor Percentual de OS no Prazo Padrão", "top10_prazo", 115),
        ("Top 3 Percentis de Desempenho – Maiores Ofensores (Ref. Últimos 3 meses)", "top3_prazo", 110),
        ("Qtde de OS e % Serviços no Prazo - Últimos 6 meses", "grafico_prazo", 80)
    ]

    for titulo, chave_img, tamanho in seccoes:
        h = doc.add_heading(titulo, level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        add_imagem_ou_texto(doc, f"[Imagem de {titulo} em anexo]", imagens[chave_img], tamanho_pct=tamanho, align='center')

    # ANEXO com a tabela
    doc.add_heading("ANEXO", level=1)
    doc.add_paragraph("Considera-se essa Tabela de Serviços Prioritários:")
    add_imagem_ou_texto(doc, "[Tabela de serviços em anexo]", IMG_TABELA, tamanho_pct=100, align='left')

    # Fonte de dados e premissas
    doc.add_heading("FONTE DE DADOS", level=1)
    doc.add_paragraph(
        "Todos os dados utilizados neste painel são provenientes da seguinte base:\n"
        "• Inova: Informações de serviços, ordens de serviço, prazo padrão e tempo padrão."
    )

    doc.add_heading("REGRAS E PREMISSAS", level=1)
    doc.add_paragraph(
        "O que foi considerado:\n"
        "• Exclusão da área de Engenharia\n"
        "• Considera o Dia do Prazo + 3 dias para os serviços do GSC que possuem apropriação manual\n"
        "• Tempo de Execução: Diferença entre início e fim da execução do serviço\n"
        "• Dias de Execução: Diferença entre Data Inclusão e Data Baixa do serviço\n"
        "• Remoção de serviços executados sem tempo de início e fim de execução"
    )

    # Salvar documento
    caminho_doc = os.path.join(pasta_saida, f"{concessionaria}_Report.docx")
    doc.save(caminho_doc)
    print(f"Documento gerado: {caminho_doc}")

# ----------- EXECUÇÃO -----------
if __name__ == "__main__":
    for conc in CONCESSIONARIAS:
        gerar_documento_report(conc, PASTA_SAIDA, IMAGENS[conc])
