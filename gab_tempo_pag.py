from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def construir_segunda_pagina(doc: Document, nome_conc: str, mes: str, ano: int, imagens: dict):
    
    # === Zerar espaçamento dos estilos ===
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(0)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    # ---------- BLOCO 1 - Indicador ----------
    h = doc.add_heading("Indicador de Serviços no Tempo Padrão", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_before = Pt(8)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img.paragraph_format.space_after = Pt(4)
    
    run_img = p_img.add_run()
    if os.path.exists(imagens["card_tempo"]):
        run_img.add_picture(imagens["card_tempo"], width=Inches(6 * 0.70))

    p_leg = doc.add_paragraph("Serviços Executados")
    p_leg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_leg.runs[0].font.size = Pt(10)
    p_leg.runs[0].italic = True
    p_leg.paragraph_format.space_after = Pt(12)

    # ---------- BLOCO 2 - Top 10 ----------
    h = doc.add_heading("Top 10 Serviços com Menor Percentual de OS no Tempo Padrão", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_before = Pt(6)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img.paragraph_format.space_after = Pt(12)
    
    run_img = p_img.add_run()
    if os.path.exists(imagens["top10_tempo"]):
        run_img.add_picture(imagens["top10_tempo"], width=Inches(6 * 1.15))

    # ---------- BLOCO 3 + BLOCO 4 LADO A LADO ----------
    h = doc.add_heading("Top 3 e Gráfico – Últimos 6 meses", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_after = Pt(8)

    tabela = doc.add_table(rows=1, cols=2)
    tabela.autofit = False

    # ----- COLUNA ESQUERDA = GRÁFICO -----
    cell1 = tabela.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = p1.add_run()

    if os.path.exists(imagens["grafico_tempo"]):
        run1.add_picture(imagens["grafico_tempo"], width=Inches(3.5))

    # ----- COLUNA DIREITA = TOP 3 COMPACTO -----
    cell2 = tabela.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run()

    caminho_top3_compacto = imagens["top3_tempo"]
    if os.path.exists(caminho_top3_compacto):
        run2.add_picture(caminho_top3_compacto, width=Inches(3.2))

    doc.add_paragraph("\n")
