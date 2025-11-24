from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def construir_primeira_pagina(doc: Document, nome_conc: str, mes: str, ano: int, imagens: dict):

    # === Zerar espaçamento dos estilos ===
    doc.styles['Heading 1'].paragraph_format.space_before = Pt(0)
    doc.styles['Heading 1'].paragraph_format.space_after = Pt(0)
    doc.styles['Normal'].paragraph_format.space_before = Pt(0)
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    # ---------- Título ----------
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)

    run = p.add_run(nome_conc)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(21)

    run2 = p.add_run(f"\nReport Cesta de Serviços – {mes} {ano}")
    run2.font.name = 'Calibri'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run2.font.size = Pt(12)

    # ---------- BLOCO 1 ----------
    h = doc.add_heading("Indicador de Serviços no Prazo Padrão", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_before = Pt(8)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img.paragraph_format.space_after = Pt(4)

    run_img = p_img.add_run()
    if os.path.exists(imagens["card_prazo"]):
        run_img.add_picture(imagens["card_prazo"], width=Inches(6 * 0.70))

    p_leg = doc.add_paragraph("Serviços Executados e Não Executados")
    p_leg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_leg.runs[0].font.size = Pt(10)
    p_leg.runs[0].italic = True
    p_leg.paragraph_format.space_after = Pt(12)

    # ---------- BLOCO 2 ----------
    h = doc.add_heading("Top 10 Serviços com Menor Percentual de OS no Prazo Padrão", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_before = Pt(6)

    p_img = doc.add_paragraph()
    p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p_img.paragraph_format.space_after = Pt(12)

    run_img = p_img.add_run()
    if os.path.exists(imagens["top10_prazo"]):
        run_img.add_picture(imagens["top10_prazo"], width=Inches(6 * 1.15))

    # ---------- BLOCO 3 / BLOCO 4 ----------
    h = doc.add_heading("Top 3 e Gráfico Qtd OS x % Serviços no Prazo – Últimos 6 meses", level=1)
    h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    h.paragraph_format.space_after = Pt(8)

    tabela = doc.add_table(rows=1, cols=2)
    tabela.autofit = False

    # Coluna esquerda
    cell1 = tabela.rows[0].cells[0]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run1 = p1.add_run()
    if os.path.exists(imagens["grafico_prazo"]):
        run1.add_picture(imagens["grafico_prazo"], width=Inches(3.5))

    # Coluna direita
    cell2 = tabela.rows[0].cells[1]
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    run2 = p2.add_run()
    if os.path.exists(imagens["top3_prazo"]):
        run2.add_picture(imagens["top3_prazo"], width=Inches(3.2))
