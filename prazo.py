from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def construir_primeira_pagina(doc: Document, nome_conc: str, mes: str, ano: int, imagens: dict):
    # ---------- Título ----------
    p = doc.add_paragraph()
    run = p.add_run(nome_conc)
    run.font.name = 'Calibri'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run.font.size = Pt(21)
    run.add_break()  # quebra de linha

    # ---------- Subtítulo ----------
    run2 = p.add_run(f"Report Cesta de Serviços – {mes} {ano}")
    run2.font.name = 'Calibri'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
    run2.font.size = Pt(12)

    # ---------- Seções principais ----------
    seccoes = [
        ("Indicador de Serviços no Prazo Padrão", "card_prazo", 70),
        ("Top 10 Serviços com Menor Percentual de OS no Prazo Padrão", "top10_prazo", 115),
        ("Top 3 Percentis de Desempenho – Maiores Ofensores (Ref. Últimos 3 meses)", "top3_prazo", 110),
        ("Qtde de OS e % Serviços no Prazo - Últimos 6 meses", "grafico_prazo", 80)
    ]
    
    textos_exp = {
        "card_prazo": "Serviços Executados e Não Executados",
        "top10_prazo": None,
        "top3_prazo": "Considerados Acima de 10 O.S",
        "grafico_prazo": "Serviços Executados e Não Executados"
    }

    distancia_entre_blocos = Pt(8)  # distância uniforme entre blocos

    for titulo, chave_img, tamanho in seccoes:
        # ---------- Título da seção ----------
        h = doc.add_heading(titulo, level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(2)

        # ---------- Imagem ----------
        p_img = doc.add_paragraph()
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        run_img = p_img.add_run()
        if imagens[chave_img] and os.path.exists(imagens[chave_img]):
            largura_inch = Inches(6) * (tamanho / 100)
            run_img.add_picture(imagens[chave_img], width=largura_inch)

        # ---------- Legenda ----------
        texto_exp = textos_exp.get(chave_img)
        p_leg = doc.add_paragraph()
        p_leg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_leg.paragraph_format.space_before = Pt(0)
        p_leg.paragraph_format.space_after = distancia_entre_blocos

        if texto_exp:
            run_leg = p_leg.add_run(texto_exp)
            run_leg.font.name = "Calibri"
            run_leg._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            run_leg.font.size = Pt(10)
            run_leg.font.italic = True
        else:
            # apenas um run vazio para garantir o espaçamento
            run_leg = p_leg.add_run("")
