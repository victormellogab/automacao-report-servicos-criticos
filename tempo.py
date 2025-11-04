from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

def construir_segunda_pagina(doc: Document, nome_conc: str, mes: str, ano: int, imagens: dict):
    # ---------- Seções principais ----------
    seccoes = [
        ("Indicador de Serviços no Tempo Padrão", "card_tempo", 70),
        ("Top 10 Serviços com Menor Percentual de OS no Tempo Padrão", "top10_tempo", 115),
        ("Top 3 Percentis de Desempenho – Maiores Ofensores (Ref. Últimos 3 meses)", "top3_tempo", 110),
        ("Qtde de OS e % Serviços no Tempo - Últimos 6 meses", "grafico_tempo", 80)
    ]
    
    textos_exp = {
        "card_tempo": "Serviços Executados",
        "top10_tempo": None,
        "top3_tempo": "Considerados Acima de 10 O.S",
        "grafico_tempo": "Serviços Executados"
    }

    distancia_entre_blocos = Pt(8)  # distância uniforme entre blocos

    for titulo, chave_img, tamanho in seccoes:
        # ---------- Título da seção ----------
        h = doc.add_heading(titulo, level=1)
        h.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        h.paragraph_format.space_before = Pt(0)
        h.paragraph_format.space_after = Pt(2)

        # ---------- Imagem ----------
        p_img = doc.add_paragraph()
        p_img.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)
        run_img = p_img.add_run()
        if imagens[chave_img] and os.path.exists(imagens[chave_img]):
            largura_inch = Inches(6) * (tamanho / 100)
            run_img.add_picture(imagens[chave_img], width=largura_inch)

        # ---------- Legenda ----------
        texto_exp = textos_exp.get(chave_img)
        p_leg = doc.add_paragraph()
        p_leg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p_leg.paragraph_format.space_before = Pt(0)
        p_leg.paragraph_format.space_after = distancia_entre_blocos

        if texto_exp:
            run_leg = p_leg.add_run(texto_exp)
            run_leg.font.name = "Calibri"
            run_leg._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
            run_leg.font.size = Pt(10)
            run_leg.font.italic = True
        else:
            # apenas um run vazio para garantir o espaçamento
            run_leg = p_leg.add_run("")
