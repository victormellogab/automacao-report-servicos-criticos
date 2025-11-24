import os
import pandas as pd
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import plotly.graph_objects as go
from prazo_pag import construir_primeira_pagina
from tempo_pag import construir_segunda_pagina
from config import CONCESSIONARIAS

# ----------- CONFIGURAÇÃO -----------
PASTA_SAIDA = r'H:\COMERCIAL\MEDIÇÃO\ARQUIVO MORTO\BERNARDO\GSC\Automação Report'
os.makedirs(PASTA_SAIDA, exist_ok=True)

CONCESSIONARIA_NOME = {
    'CAAN': 'Concessionária Águas das Agulhas Negras',
    'CAC': 'Concessionária Águas da Condessa',
    'CAI': 'Concessionária Águas do Imperador',
    'CAIZ': 'Concessionária Águas da Imperatriz',
    'CAJ': 'Concessionária Águas de Juturnaíba',
    'CAJA': 'Concessionária Águas de Jahu',
    'CAN': 'Concessionária Águas de Niterói',
    'CANF': 'Concessionária Águas de Nova Friburgo',
    'CAP': 'Concessionária Águas do Paraíba',
    'CAPAM': 'Concessionária Águas de Pará de Minas',
    'CAPY': 'Concessionária Águas de Paraty',
    'CAV': 'Concessionária Águas de Votorantim',
}

MES = "Outubro"
ANO = 2025

IMAGENS = {}
IMG_TABELA = os.path.join(PASTA_SAIDA, "Tabela_Prioritarios.png")
CAMINHO_TABELA = r"C:\dados-report\Cesta de Serviços.xlsx"

for conc in CONCESSIONARIAS:
    IMAGENS[conc] = {
        "card_prazo": os.path.join(PASTA_SAIDA, f"{conc}_Prazo_Cards.png"),
        "top10_prazo": os.path.join(PASTA_SAIDA, f"{conc}_Prazo_Top10.png"),
        "top3_prazo": os.path.join(PASTA_SAIDA, f"{conc}_Prazo_Top3.png"),
        "grafico_prazo": os.path.join(PASTA_SAIDA, f"{conc}_Prazo_Grafico_6meses.png"),

        "card_tempo": os.path.join(PASTA_SAIDA, f"{conc}_Tempo_Cards.png"),
        "top10_tempo": os.path.join(PASTA_SAIDA, f"{conc}_Tempo_Top10.png"),
        "top3_tempo": os.path.join(PASTA_SAIDA, f"{conc}_Tempo_Top3.png"),
        "grafico_tempo": os.path.join(PASTA_SAIDA, f"{conc}_Tempo_Grafico_6meses.png"),

        "tabela_prioritarios": IMG_TABELA  # ainda usar a mesma tabela
    }

# ----------- FUNÇÕES UTILS -----------
def gerar_imagem_tabela(excel_path, caminho_saida):
    df = pd.read_excel(excel_path)
    max_lens = [max(df[col].astype(str).map(len).max(), len(col)) for col in df.columns]
    max_lens[0] = int(max_lens[0] * 0.16)
    total_len = sum(max_lens)
    column_widths = [length / total_len for length in max_lens]

    fig = go.Figure(data=[go.Table(
        columnwidth=column_widths,
        header=dict(values=list(df.columns), fill_color='darkblue', font=dict(color='white', size=14), align='center'),
        cells=dict(values=[df[col] for col in df.columns], fill_color='white', font=dict(color='black', size=12),
                   align=['left' if col=='Serviços' else 'center' for col in df.columns])
    )])
    altura = 21 + 21 * len(df)
    fig.update_layout(margin=dict(l=0, r=0, t=0, b=0), width=None, height=altura)
    fig.write_image(caminho_saida, scale=3)

def add_imagem_ou_texto(doc, texto, caminho_imagem, tamanho_pct=100, align='center'):
    p = doc.add_paragraph()
    p.alignment = {'center': WD_PARAGRAPH_ALIGNMENT.CENTER, 'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
                   'right': WD_PARAGRAPH_ALIGNMENT.RIGHT}.get(align, WD_PARAGRAPH_ALIGNMENT.CENTER)
    run = p.add_run()
    if caminho_imagem and os.path.exists(caminho_imagem):
        run.add_picture(caminho_imagem, width=Inches(6)*(tamanho_pct/100))
    else:
        run.add_text(texto)

# ----------- FUNÇÃO DE CRIAÇÃO DO DOCUMENTO -----------
def gerar_documento_report(concessionaria, pasta_saida, imagens):
    gerar_imagem_tabela(CAMINHO_TABELA, IMG_TABELA)
    
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # Fonte padrão
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')

    # Configura estilos de títulos
    for level in range(1, 10):
        try:
            s = doc.styles[f'Heading {level}']
            s.font.name = 'Calibri'
            s.font.size = Pt(12)
            s.font.bold = False
            s.font.color.rgb = RGBColor(0,0,0)
            s.element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri')
        except KeyError:
            pass

    # Nome completo da concessionária
    nome_conc = CONCESSIONARIA_NOME.get(concessionaria, concessionaria)

    # ---------- CONSTRUÇÃO DA PRIMEIRA PÁGINA ----------
    construir_primeira_pagina(doc, nome_conc, MES, ANO, imagens)
    
    # ---------- QUEBRA DE PÁGINA ----------
    doc.add_page_break()

    # ---------- CONSTRUÇÃO DA SEGUNDA PÁGINA ----------
    construir_segunda_pagina(doc, nome_conc, MES, ANO, imagens)
    
    # ---------- QUEBRA DE PÁGINA ----------
    doc.add_page_break()

    # ---------- ANEXO ----------
    doc.add_heading("ANEXO", level=1)
    doc.add_paragraph("Considera-se essa Tabela de Serviços Prioritários:")
    add_imagem_ou_texto(doc, "[Tabela de serviços em anexo]", IMG_TABELA, tamanho_pct=100, align='left')

    # Fonte de dados e premissas
    doc.add_heading("FONTE DE DADOS", level=1)
    doc.add_paragraph(
        "Todos os dados utilizados neste painel são provenientes da seguinte base:\n"
        "• Inova: Informações de serviços, ordens de serviço, prazo padrão e tempo padrão."
    )

    doc.add_heading("REGRAS E PREMISSAS", level=1)
    doc.add_paragraph(
        "O que foi considerado:\n"
        "• Exclusão da área de Engenharia\n"
        "• Considera o Dia do Prazo + 3 dias para os serviços do GSC que possuem apropriação manual\n"
        "• Tempo de Execução: Diferença entre início e fim da execução do serviço\n"
        "• Dias de Execução: Diferença entre Data Inclusão e Data Baixa do serviço\n"
        "• Remoção de serviços executados sem tempo de início e fim de execução"
    )

    caminho_doc = os.path.join(pasta_saida, f"{concessionaria}_Report.docx")
    doc.save(caminho_doc)
    print(f"Documento gerado: {caminho_doc}")

# ----------- EXECUÇÃO -----------
if __name__ == "__main__":
    for conc in CONCESSIONARIAS:
        gerar_documento_report(conc, PASTA_SAIDA, IMAGENS[conc])
